from docx import Document
import zipfile
import os
import shutil

# Read template
doc = Document('dsa_template.docx')

print(f'Total paragraphs: {len(doc.paragraphs)}')
print(f'Total tables: {len(doc.tables)}')
print('\n--- Document Structure ---')

for i, p in enumerate(doc.paragraphs[:30]):
    text = p.text.strip()
    if text:
        print(f'{i}: {text[:100]}')

# Extract images
print('\n--- Extracting Images ---')
with zipfile.ZipFile('dsa_template.docx', 'r') as z:
    media_files = [f for f in z.namelist() if 'media' in f]
    print(f'Found {len(media_files)} images')
    
    # Extract each image
    for media in media_files:
        filename = os.path.basename(media)
        z.extract(media, 'temp_extract')
        src = os.path.join('temp_extract', media)
        dst = filename
        shutil.copy(src, dst)
        print(f'Extracted: {filename}')
    
    # Cleanup
    if os.path.exists('temp_extract'):
        shutil.rmtree('temp_extract')

print('\nDone!')
